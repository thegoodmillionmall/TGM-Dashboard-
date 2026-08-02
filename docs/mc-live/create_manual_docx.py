from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "คู่มือ-MC-Live.docx"
SCREENSHOTS = BASE / "screenshots"


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border_top(cell, color="1A2A3A", size="12"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:color"), color)
    borders.append(top)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 20, True, "0B2239")
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        set_font(r, 11, False, "6B7A90")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 15 if level == 1 else 12, True, "0B2239")
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.add_run("• " + text)
    set_font(r, 10.5)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    shade(cell, "E8F6F5")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_font(r, 10.5, False, "0B2239")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "DFF3F3")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 9.5, True, "0B2239")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_font(r, 9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_image(doc, filename, caption):
    path = SCREENSHOTS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    r = p.add_run(caption)
    set_font(r, 10.5, True, "0B2239")
    doc.add_picture(str(path), width=Inches(6.6))
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)

styles = doc.styles
styles["Normal"].font.name = "Tahoma"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
styles["Normal"].font.size = Pt(10.5)

add_title(doc, "คู่มือการใช้งาน MC Live", "The Good Million | สำหรับ MC, หัวหน้า MC และผู้บริหาร")
add_note(doc, "ระบบนี้เริ่มใช้จริงวันที่ 1 สิงหาคม 2026 ใช้สำหรับบันทึกยอดไลฟ์ ตรวจหลักฐาน และอนุมัติยอดจริงรายเดือน")

add_heading(doc, "ภาพรวม Flow ระบบ")
add_table(
    doc,
    ["ลำดับ", "ผู้เกี่ยวข้อง", "สิ่งที่ทำ", "ผลลัพธ์"],
    [
        ["1", "MC", "กรอกยอดขาย เวลาเริ่ม-สิ้นสุด ออเดอร์ Ads/Coins และแนบหลักฐาน", "รายการอยู่ในสถานะรอหัวหน้าเช็ค"],
        ["2", "หัวหน้า MC", "เปิดดูหลักฐานแบบ preview ตรวจความครบถ้วน และอนุมัติหรือส่งกลับแก้ไข", "รายการที่ถูกต้องถูก mark ว่าเช็คแล้ว"],
        ["3", "ผู้บริหาร", "ดูภาพรวมรายเดือน รายวัน รายคน และอนุมัติยอดจริงสิ้นเดือน", "ยอดเดือนนั้นถูกล็อกเป็นยอดจริง"],
        ["4", "กรณีแก้ไข", "ผู้บริหารเปิดเดือนกลับมาแก้ไขได้", "ทีมแก้ไขแล้วส่งตรวจใหม่"],
    ],
    [0.5, 1.2, 3.1, 2.0],
)

add_heading(doc, "1. คู่มือสำหรับ MC")
for text in [
    "เข้าเมนู MC Live แล้วเลือกแท็บ “กรอกของฉัน”",
    "เลือกบริษัทก่อน โดยมี Nola และ TGM ถ้าเป็น Nola ระบบจะให้กรอกเฉพาะ TikTok",
    "เลือก Platform, ประเภทกล้อง OBS/มือถือ, วันที่, เวลาเริ่มต้น, เวลาสิ้นสุด",
    "กรอกยอดขาย จำนวนออเดอร์ ค่า Ads และ Coins ตามข้อมูลจริง",
    "ระบบจะคำนวณจำนวนชั่วโมงไลฟ์และยอดขายต่อชั่วโมงให้อัตโนมัติ",
    "ถ้าเลือกมือถือ ต้องแนบ 3 ภาพ: ภาพหน้าจอที่ไลฟ์, หน้ายอดขาย, หน้าจบไลฟ์",
    "ถ้าเลือก OBS แนบ 1 ภาพหน้าจอหลักฐานก็พอ",
]:
    add_bullet(doc, text)
add_image(doc, "02-mc-entry.png", "ตัวอย่างหน้ากรอกของ MC")

add_heading(doc, "2. คู่มือสำหรับหัวหน้า MC")
for text in [
    "เข้าแท็บ “หัวหน้าเช็ค” เพื่อดูรายการที่ทีมส่งเข้ามา",
    "กดรายการหรือปุ่มหลักฐานเพื่อเปิด preview รูปใน popup",
    "ตรวจว่าแพลตฟอร์ม ยอดขาย เวลา และเอกสารตรงกัน",
    "ถ้าถูกต้องให้กดเช็คแล้ว ระบบจะแสดงสถานะว่าเช็คแล้ว",
    "ถ้าไม่ถูกต้องให้ส่งกลับแก้ไข พร้อมระบุหมายเหตุให้ MC แก้",
]:
    add_bullet(doc, text)
add_image(doc, "03-lead-review-list.png", "ตัวอย่างหน้าหัวหน้าเช็คหลักฐาน")

add_heading(doc, "3. คู่มือสำหรับผู้บริหาร")
for text in [
    "เข้าแท็บ “ภาพรวม” เพื่อดูยอดรวมรายเดือน จำนวนไลฟ์ ชั่วโมงไลฟ์ ออเดอร์ ค่า Ads และสถานะเช็คหลักฐาน",
    "ตรวจ performance รายคน เพื่อดูยอดขายรวม ชั่วโมงรวม ออเดอร์ และค่าเฉลี่ยต่อชั่วโมง",
    "ก่อนปิดเดือน ให้เข้าแท็บ “อนุมัติรายเดือน”",
    "ถ้าข้อมูลถูกต้อง ให้กดอนุมัติยอดจริงเดือนนี้ ยอดเดือนนั้นจะถือเป็นยอดจริง",
    "ถ้ามีรายการผิด ให้เปิดเดือนกลับมาแก้ไขก่อน แล้วให้ทีมส่งตรวจใหม่",
]:
    add_bullet(doc, text)
add_image(doc, "01-executive-overview.png", "ตัวอย่างหน้าภาพรวมผู้บริหาร")
add_image(doc, "04-executive-month-approve.png", "ตัวอย่างหน้าอนุมัติรายเดือน")

add_heading(doc, "4. การแก้ไขรายการ")
for text in [
    "แท็บแก้ไขตารางใช้สำหรับผู้มีสิทธิ์แก้รายการจำนวนมาก",
    "ควรใช้เมื่อมีข้อมูลหลายรายการต้องปรับพร้อมกัน",
    "รายการที่ผู้บริหารอนุมัติรายเดือนไปแล้ว ต้องเปิดเดือนกลับมาแก้ไขก่อน",
]:
    add_bullet(doc, text)
add_image(doc, "05-edit-table.png", "ตัวอย่างหน้าแก้ไขตาราง")

add_heading(doc, "5. สิทธิ์ผู้ใช้งาน")
add_table(
    doc,
    ["บทบาท", "สิทธิ์หลัก", "ข้อจำกัด"],
    [
        ["MC", "กรอกและดูรายการของตัวเอง", "แก้ได้เฉพาะรายการตัวเอง และแก้ไม่ได้หลังอนุมัติรายเดือน"],
        ["หัวหน้า MC", "ดูรายการทีม ตรวจหลักฐาน อนุมัติ/ส่งกลับแก้ไข", "ไม่ควรอนุมัติรายการที่หลักฐานไม่ครบ"],
        ["ผู้บริหาร/Admin", "ดูภาพรวม แก้ไข เปิดเดือน และอนุมัติยอดจริงรายเดือน", "เมื่ออนุมัติแล้วถือเป็นยอดจริงของเดือน"],
    ],
    [1.2, 3.0, 2.6],
)

add_heading(doc, "6. Checklist ก่อนใช้งานจริง")
for text in [
    "ลบข้อมูลตัวอย่างออกจากระบบแล้วก่อนเปิดใช้งานจริง",
    "สร้าง user จริงให้น้อง MC ทุกคน",
    "แจ้งรหัสผ่านชั่วคราวและให้เปลี่ยนรหัสเมื่อเข้าใช้งานครั้งแรก",
    "ทดลองให้ MC 1 คนกรอกจริง 1 รายการ และให้หัวหน้าเช็ค 1 รายการ",
    "ผู้บริหารทดลองเปิดหน้าอนุมัติรายเดือนก่อนสิ้นเดือน",
]:
    add_bullet(doc, text)

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "ภาคผนวก: คำอธิบายตัวเลข KPI")
add_table(
    doc,
    ["KPI", "ความหมาย"],
    [
        ["ยอดขายรวม", "ยอดขายจากรายการไลฟ์ทั้งหมดในช่วงวันที่ที่เลือก"],
        ["จำนวนไลฟ์", "จำนวนรายการไลฟ์ที่ถูกบันทึก"],
        ["ชั่วโมงไลฟ์รวม", "ผลรวมเวลาสิ้นสุดลบเวลาเริ่มต้นของทุกรายการ"],
        ["ออเดอร์รวม", "จำนวนออเดอร์ที่ MC กรอกในแต่ละไลฟ์รวมกัน"],
        ["ค่า Ads", "ค่าโฆษณาที่ใช้กับไลฟ์ในช่วงนั้น"],
        ["ยอดขายต่อชั่วโมง", "ยอดขายรวมหารด้วยชั่วโมงไลฟ์รวม"],
        ["เช็คหลักฐาน", "จำนวนรายการที่หัวหน้า MC ตรวจหลักฐานแล้วเทียบกับรายการทั้งหมด"],
    ],
    [1.8, 5.0],
)

doc.save(OUT)
print(OUT)
